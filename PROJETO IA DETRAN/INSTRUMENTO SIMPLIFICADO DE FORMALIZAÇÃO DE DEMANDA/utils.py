import io
import os
import sys
import inspect
import functools
from pathlib import Path
from enum import IntEnum, auto
from docx import Document


def with_file_line(func):
    @functools.wraps(func)
    def wrapper(*msg, **kwargs):
        frame = inspect.currentframe().f_back
        line_no = kwargs.get("line_no", None)
        if line_no is None:
            line_no = frame.f_lineno
        file_name = kwargs.get("file_name", None)
        if file_name is None:
            file_name = Path(inspect.getframeinfo(frame).filename).relative_to(os.getcwd())
        return func(*msg, file_name=file_name, line_no=line_no)

    return wrapper


class bcolors:
    # fmt: off
    HEADER    = "\033[95m"
    OKBLUE    = "\033[94m"
    OKCYAN    = "\033[96m"
    OKGREEN   = "\033[92m"
    WARNING   = "\033[93m"
    FAIL      = "\033[91m"
    ENDC      = "\033[0m"
    BOLD      = "\033[1m"
    UNDERLINE = "\033[4m"
    # fmt: on


@with_file_line
def TODO(*msg, file_name="", line_no=""):
    sys.stdout.flush()
    print(bcolors.WARNING, f"{file_name}:{line_no}: TODO:", *msg, bcolors.ENDC, file=sys.stderr, flush=True)
    os.abort()


@with_file_line
def UNREACHABLE(*msg, file_name="", line_no=""):
    sys.stdout.flush()
    print(bcolors.FAIL, f"{file_name}:{line_no}: UNREACHABLE:", *msg, bcolors.ENDC, file=sys.stderr, flush=True)
    os.abort()


class LogLevel(IntEnum):
    # fmt: off
    INFO  = auto()
    WARN  = auto()
    ERROR = auto()
    NONE  = auto()
    # fmt: on


# Valor Padrao
minLogLevel = LogLevel.INFO


def setLogLevel(level: LogLevel):
    global minLogLevel
    minLogLevel = level


def getLogLevel():
    global minLogLevel
    return minLogLevel


# class LineNo:
#     def __str__(self):
#         return str(inspect.currentframe().f_back.f_lineno)


# __line__ = LineNo()


def log(level: LogLevel, *msg, **kwargs):
    if level is None:
        return

    if level < getLogLevel():
        return

    output = kwargs.get("file", None)

    frame = inspect.currentframe().f_back
    line_no = frame.f_lineno
    file_name = Path(inspect.getframeinfo(frame).filename).relative_to(os.getcwd())

    if output is sys.stdin:
        UNREACHABLE("Nao se pode escrever no stdin", file_name=file_name, line_no=line_no)

    match level:
        case LogLevel.INFO:
            print(f"{bcolors.OKGREEN}INFO:{bcolors.ENDC} ", end="", file=output)
        case LogLevel.WARN:
            print(f"{bcolors.WARNING}WARN:{bcolors.ENDC} ", end="", file=output)
        case LogLevel.ERROR:
            print(f"{bcolors.FAIL}ERROR:{bcolors.ENDC} ", end="", file=output)
        case LogLevel.NONE:
            return
        case _:
            UNREACHABLE("Invalid Log Level", file_name=file_name, line_no=line_no)

    print(*msg, **kwargs)


def extract_gentext_from_dfd(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))

        obj_sint = None
        obj_sint_fnd = False

        just = None
        just_fnd = False

        for p in doc.paragraphs:
            text = p.text.strip()
            if len(text):
                if obj_sint is not None and just is not None:
                    break
                if obj_sint_fnd:
                    obj_sint = text
                    obj_sint_fnd = False
                if just_fnd:
                    just = text
                    just_fnd = False
                if text == "8  - Objetivo da Contratação/Resultados Esperados:":
                    obj_sint_fnd = True
                    continue
                if text == "10 - Justificativa para aquisição:":
                    just_fnd = True
                    continue
        return obj_sint, just
    except Exception as e:
        # Se houver um erro, exibe a mensagem detalhada e retorna None para tratamento na interface
        log(LogLevel.ERROR, f"❌ Erro ao ler o documento: '{e}'")
        return None


def extract_items_from_dfd(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))

        for table in doc.tables:
            header_row_index = -1

            # Encontra o índice da linha de cabeçalho
            for i, row in enumerate(table.rows):
                row_text = " ".join([cell.text.strip().upper() for cell in row.cells])
                if "SIAGO/TCE" in row_text and "UN." in row_text and "QTDE" in row_text and "ESPECIFICAÇÃO DO PRODUTO" in row_text:
                    header_row_index = i
                    break

            if header_row_index != -1:
                header_row = table.rows[header_row_index]
                header_texts = [cell.text.strip().upper() for cell in header_row.cells]

                catmat_idx = -1
                desc_idx = -1
                un_idx = -1
                qtd_idx = -1

                for i, text in enumerate(header_texts):
                    if "SIAGO/TCE" in text:
                        catmat_idx = i
                    elif "ESPECIFICAÇÃO" in text or "PRODUTO" in text:
                        desc_idx = i
                    elif "UN." in text:
                        un_idx = i
                    elif "QTDE" in text:
                        qtd_idx = i

                if desc_idx != -1 and un_idx != -1 and qtd_idx != -1:
                    items_list = []

                    # Itera sobre as linhas de dados a partir da linha seguinte ao cabeçalho
                    for row in table.rows[header_row_index + 1 :]:
                        cells = row.cells
                        if len(cells) > max(catmat_idx, desc_idx, un_idx, qtd_idx):
                            item_data = {
                                "catmat": cells[catmat_idx].text.strip() if catmat_idx != -1 else "",
                                "descricao": cells[desc_idx].text.strip(),
                                "unidade": cells[un_idx].text.strip(),
                                "qtd": cells[qtd_idx].text.strip(),
                                "valor_unitario": "",
                            }
                            # Adiciona o item somente se a descrição não estiver vazia
                            if item_data["descricao"].strip():
                                items_list.append(item_data)
                    return items_list
        return []
    except Exception as e:
        # Se houver um erro, exibe a mensagem detalhada e retorna None para tratamento na interface
        log(LogLevel.ERROR, f"❌ Erro ao ler o documento: '{e}'")
        return None


# Testes
def main():
    log(LogLevel.INFO, "stdin", file=sys.stdin)
    TODO("Teste TODO")
    UNREACHABLE("Teste UNREACHABLE")


if __name__ == "__main__":
    main()
